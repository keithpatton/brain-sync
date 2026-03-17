"""Convert .docx files to markdown with comment extraction.

Primary use case: Google Docs exports where comments (engineering feedback)
are only preserved in .docx format. Supports two workflows:

1. Hybrid: Google markdown export for body + .docx only for comments
2. Full: .docx for both body and comments (lower fidelity body conversion)
"""

from __future__ import annotations

import logging
import re
from io import BytesIO
from pathlib import Path

from docx import Document

log = logging.getLogger(__name__)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _collect_text_between(start_elem, end_elem) -> str:
    """Collect all w:t text between two XML elements in the document body."""
    texts: list[str] = []
    collecting = False
    body = start_elem.getparent()

    for elem in body.iter():
        if elem is start_elem:
            collecting = True
            continue
        if elem is end_elem:
            break
        if collecting and elem.tag == f"{{{_W_NS}}}t" and elem.text:
            texts.append(elem.text)

    return " ".join(texts).strip()


def _extract_comments_from_doc(doc: Document) -> list[dict]:  # pyright: ignore[reportGeneralTypeIssues]
    """Extract comments with metadata and annotated text from a Document."""
    from lxml import etree  # pyright: ignore[reportAttributeAccessIssue]

    comments_part = None
    for rel in doc.part.rels.values():
        if "comments" in rel.reltype:
            comments_part = rel.target_part
            break

    if comments_part is None:
        return []

    comments_xml = etree.fromstring(comments_part.blob)
    comments_by_id: dict[str, dict] = {}

    for comment_elem in comments_xml.findall(f"{{{_W_NS}}}comment"):
        cid = comment_elem.get(f"{{{_W_NS}}}id")
        if cid is None:
            continue

        author = comment_elem.get(f"{{{_W_NS}}}author", "Unknown")
        date_str = comment_elem.get(f"{{{_W_NS}}}date", "")

        text_parts = []
        for para in comment_elem.findall(f".//{{{_W_NS}}}p"):
            para_texts = []
            for t in para.findall(f".//{{{_W_NS}}}t"):
                if t.text:
                    para_texts.append(t.text)
            if para_texts:
                text_parts.append("".join(para_texts))

        date_display = date_str[:10] if date_str and len(date_str) >= 10 else date_str
        comments_by_id[cid] = {
            "author": author,
            "date": date_display,
            "text": "\n".join(text_parts),
            "annotated_text": "",
        }

    if not comments_by_id:
        return []

    body = doc.element.body
    for cid, info in comments_by_id.items():
        starts = body.findall(f".//{{{_W_NS}}}commentRangeStart[@{{{_W_NS}}}id='{cid}']")
        ends = body.findall(f".//{{{_W_NS}}}commentRangeEnd[@{{{_W_NS}}}id='{cid}']")
        if starts and ends:
            annotated = _collect_text_between(starts[0], ends[0])
            info["annotated_text"] = annotated[:200]

    return sorted(comments_by_id.values(), key=lambda c: (c.get("date", ""), c.get("author", "")))


def _format_comments_markdown(comments: list[dict]) -> str:
    parts: list[str] = []
    for c in comments:
        author = c.get("author", "Unknown")
        date = c.get("date", "")
        text = c.get("text", "")
        annotated = c.get("annotated_text", "")

        header = f"### {author}"
        if date:
            header += f" ({date})"
        parts.append(header)

        if annotated:
            parts.append(f'**On:** "{annotated}"')
            parts.append("")

        parts.append(text)
        parts.append("")

    return "\n".join(parts).rstrip()


def extract_comments(docx_path: Path) -> str | None:
    """Extract comments from a .docx file as formatted markdown."""
    doc = Document(str(docx_path))
    comments = _extract_comments_from_doc(doc)
    if not comments:
        return None
    return _format_comments_markdown(comments)


def extract_comments_from_bytes(data: bytes) -> str | None:
    """Extract comments from .docx bytes as formatted markdown."""
    doc = Document(BytesIO(data))
    comments = _extract_comments_from_doc(doc)
    if not comments:
        return None
    return _format_comments_markdown(comments)


def _para_to_markdown(para) -> str:
    """Convert a single docx paragraph to markdown."""
    style_name = (para.style.name or "").lower() if para.style else ""

    heading_match = re.match(r"heading\s*(\d+)", style_name)
    prefix = ""
    if heading_match:
        level = int(heading_match.group(1))
        prefix = "#" * min(level, 6) + " "

    if para._element.find(f".//{{{_W_NS}}}numPr") is not None:
        ilvl_elem = para._element.find(f".//{{{_W_NS}}}ilvl")
        indent = int(ilvl_elem.get(f"{{{_W_NS}}}val", "0")) if ilvl_elem is not None else 0
        prefix = "  " * indent + "- "

    parts: list[str] = []
    for run in para.runs:
        text = run.text or ""
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)

    return prefix + "".join(parts)


def _table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip().replace("|", "\\|") for cell in row.cells])

    if not rows:
        return ""

    lines: list[str] = []
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    for row in rows[1:]:
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")
    return "\n".join(lines)


def docx_to_markdown(docx_path: Path) -> str:
    """Convert a .docx file to markdown with comments appended."""
    doc = Document(str(docx_path))
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            for para in doc.paragraphs:
                if para._element is element:
                    parts.append(_para_to_markdown(para))
                    break
        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    parts.append("")
                    parts.append(_table_to_markdown(table))
                    parts.append("")
                    break

    body = "\n".join(parts).strip() + "\n"

    comments = _extract_comments_from_doc(doc)
    if comments:
        comments_md = _format_comments_markdown(comments)
        body = body.rstrip("\n") + "\n\n---\n\n## Comments\n\n" + comments_md + "\n"

    return body


def append_comments_to_markdown(md_path: Path, docx_path: Path) -> bool:
    """Read .md, extract comments from .docx, append ## Comments section."""
    comments_md = extract_comments(docx_path)
    if comments_md is None:
        log.info("No comments found in %s", docx_path.name)
        return False

    md_content = md_path.read_text(encoding="utf-8")
    md_content = re.sub(r"\n---\n\n## Comments\n.*", "", md_content, flags=re.DOTALL)
    md_content = re.sub(r"\n## Comments\n.*", "", md_content, flags=re.DOTALL)

    new_content = md_content.rstrip("\n") + "\n\n---\n\n## Comments\n\n" + comments_md + "\n"
    md_path.write_text(new_content, encoding="utf-8")

    log.info("Appended comments from %s to %s", docx_path.name, md_path.name)
    return True


__all__ = [
    "_W_NS",
    "append_comments_to_markdown",
    "docx_to_markdown",
    "extract_comments",
    "extract_comments_from_bytes",
]
