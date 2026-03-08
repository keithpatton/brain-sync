"""Tests for the .docx comment extraction and conversion module."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

from brain_sync.docx_converter import (
    append_comments_to_markdown,
    docx_to_markdown,
    extract_comments,
    extract_comments_from_bytes,
)


# ---------------------------------------------------------------------------
# Helpers for creating .docx fixtures with comments
# ---------------------------------------------------------------------------

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def _make_simple_docx(tmp_path: Path, name: str = "test.docx") -> Path:
    """Create a simple .docx with headings, paragraphs, bold/italic."""
    doc = Document()
    doc.add_heading("Project Overview", level=1)
    p = doc.add_paragraph()
    p.add_run("This is a ").bold = False
    bold_run = p.add_run("bold statement")
    bold_run.bold = True
    p.add_run(" and ")
    italic_run = p.add_run("italic text")
    italic_run.italic = True
    p.add_run(".")
    doc.add_heading("Details", level=2)
    doc.add_paragraph("Some detail content here.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def _add_comments_to_docx(
    docx_path: Path,
    comments: list[dict],
) -> None:
    """Add multiple comments to an existing .docx file.

    Each comment dict has keys: id, author, date, text, para_idx.
    This manipulates OOXML directly since python-docx has no write API
    for comments.
    """
    doc = Document(str(docx_path))

    # Create comments XML
    comments_xml = etree.Element(
        f"{{{_W_NS}}}comments",
        nsmap={"w": _W_NS},
    )

    for c in comments:
        comment_elem = etree.SubElement(comments_xml, f"{{{_W_NS}}}comment")
        comment_elem.set(f"{{{_W_NS}}}id", c["id"])
        comment_elem.set(f"{{{_W_NS}}}author", c["author"])
        comment_elem.set(f"{{{_W_NS}}}date", c["date"])
        cp = etree.SubElement(comment_elem, f"{{{_W_NS}}}p")
        cr = etree.SubElement(cp, f"{{{_W_NS}}}r")
        ct_elem = etree.SubElement(cr, f"{{{_W_NS}}}t")
        ct_elem.text = c["text"]

    # Create comments part
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    comments_blob = etree.tostring(comments_xml, xml_declaration=True, encoding="UTF-8")
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
    part_name = PackURI("/word/comments.xml")
    comments_part = Part(part_name, content_type, comments_blob, doc.part.package)
    doc.part.relate_to(
        comments_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    )

    # Add commentRangeStart/End markers in document body
    body = doc.element.body
    paras = body.findall(f"{{{_W_NS}}}p")
    for c in comments:
        para_idx = c.get("para_idx", 0)
        if para_idx < len(paras):
            target_para = paras[para_idx]
            range_start = etree.Element(f"{{{_W_NS}}}commentRangeStart")
            range_start.set(f"{{{_W_NS}}}id", c["id"])
            target_para.insert(0, range_start)
            range_end = etree.Element(f"{{{_W_NS}}}commentRangeEnd")
            range_end.set(f"{{{_W_NS}}}id", c["id"])
            target_para.append(range_end)

    doc.save(str(docx_path))


def _make_docx_with_comments(tmp_path: Path, name: str = "commented.docx") -> Path:
    """Create a .docx with body text and comments."""
    path = _make_simple_docx(tmp_path, name)

    _add_comments_to_docx(path, [
        {
            "id": "1",
            "author": "Jane Smith",
            "date": "2026-03-05T10:30:00Z",
            "text": "Should we consider canary deployments instead?",
            "para_idx": 1,
        },
        {
            "id": "2",
            "author": "Bob Jones",
            "date": "2026-03-06T14:00:00Z",
            "text": "This section needs more detail.",
            "para_idx": 3,
        },
    ])
    return path


def _make_docx_with_table(tmp_path: Path) -> Path:
    """Create a .docx with a table."""
    doc = Document()
    doc.add_heading("Data Report", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "100"
    table.cell(2, 0).text = "Beta"
    table.cell(2, 1).text = "200"
    doc.add_paragraph("End of report.")
    path = tmp_path / "table.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Tests: extract_comments
# ---------------------------------------------------------------------------

class TestExtractComments:
    def test_no_comments_returns_none(self, tmp_path):
        """A .docx with no comments returns None."""
        path = _make_simple_docx(tmp_path)
        result = extract_comments(path)
        assert result is None

    def test_extracts_comment_author_and_text(self, tmp_path):
        """Comments include author and comment text."""
        path = _make_docx_with_comments(tmp_path)
        result = extract_comments(path)
        assert result is not None
        assert "Jane Smith" in result
        assert "canary deployments" in result
        assert "Bob Jones" in result
        assert "needs more detail" in result

    def test_extracts_comment_date(self, tmp_path):
        """Comments include formatted date."""
        path = _make_docx_with_comments(tmp_path)
        result = extract_comments(path)
        assert "2026-03-05" in result
        assert "2026-03-06" in result

    def test_comments_sorted_by_date(self, tmp_path):
        """Comments are sorted chronologically."""
        path = _make_docx_with_comments(tmp_path)
        result = extract_comments(path)
        jane_pos = result.index("Jane Smith")
        bob_pos = result.index("Bob Jones")
        assert jane_pos < bob_pos  # Jane's comment is earlier

    def test_extracts_annotated_text(self, tmp_path):
        """Comments include the annotated text excerpt."""
        path = _make_docx_with_comments(tmp_path)
        result = extract_comments(path)
        # The annotated text should appear with "On:" prefix
        assert "**On:**" in result


class TestExtractCommentsFromBytes:
    def test_from_bytes(self, tmp_path):
        """extract_comments_from_bytes works with raw bytes."""
        path = _make_docx_with_comments(tmp_path)
        data = path.read_bytes()
        result = extract_comments_from_bytes(data)
        assert result is not None
        assert "Jane Smith" in result

    def test_no_comments_from_bytes(self, tmp_path):
        """Returns None for .docx bytes with no comments."""
        path = _make_simple_docx(tmp_path)
        data = path.read_bytes()
        result = extract_comments_from_bytes(data)
        assert result is None


# ---------------------------------------------------------------------------
# Tests: docx_to_markdown
# ---------------------------------------------------------------------------

class TestDocxToMarkdown:
    def test_converts_headings(self, tmp_path):
        """Heading styles are converted to # prefixes."""
        path = _make_simple_docx(tmp_path)
        md = docx_to_markdown(path)
        assert "# Project Overview" in md
        assert "## Details" in md

    def test_converts_bold_italic(self, tmp_path):
        """Bold and italic formatting is preserved."""
        path = _make_simple_docx(tmp_path)
        md = docx_to_markdown(path)
        assert "**bold statement**" in md
        assert "*italic text*" in md

    def test_converts_tables(self, tmp_path):
        """Tables are converted to markdown tables."""
        path = _make_docx_with_table(tmp_path)
        md = docx_to_markdown(path)
        assert "| Name | Value |" in md
        assert "| Alpha | 100 |" in md
        assert "| Beta | 200 |" in md
        # Should have separator row
        assert "| --- | --- |" in md

    def test_appends_comments_section(self, tmp_path):
        """When comments exist, they're appended as ## Comments."""
        path = _make_docx_with_comments(tmp_path)
        md = docx_to_markdown(path)
        assert "## Comments" in md
        assert "Jane Smith" in md

    def test_no_comments_section_when_none(self, tmp_path):
        """No ## Comments section when document has no comments."""
        path = _make_simple_docx(tmp_path)
        md = docx_to_markdown(path)
        assert "## Comments" not in md

    def test_empty_document(self, tmp_path):
        """An empty document produces minimal output."""
        doc = Document()
        path = tmp_path / "empty.docx"
        doc.save(str(path))
        md = docx_to_markdown(path)
        assert isinstance(md, str)

    def test_plain_paragraph(self, tmp_path):
        """Plain paragraphs are preserved."""
        path = _make_simple_docx(tmp_path)
        md = docx_to_markdown(path)
        assert "Some detail content here." in md


# ---------------------------------------------------------------------------
# Tests: append_comments_to_markdown
# ---------------------------------------------------------------------------

class TestAppendCommentsToMarkdown:
    def test_appends_comments_to_clean_md(self, tmp_path):
        """Comments are appended to an .md file without existing comments."""
        md_path = tmp_path / "doc.md"
        md_path.write_text("# My Document\n\nSome content.\n", encoding="utf-8")
        docx_path = _make_docx_with_comments(tmp_path)

        result = append_comments_to_markdown(md_path, docx_path)
        assert result is True

        content = md_path.read_text(encoding="utf-8")
        assert "# My Document" in content
        assert "Some content." in content
        assert "## Comments" in content
        assert "Jane Smith" in content
        assert "Bob Jones" in content

    def test_replaces_existing_comments_section(self, tmp_path):
        """Existing ## Comments section is replaced (idempotent)."""
        md_path = tmp_path / "doc.md"
        md_path.write_text(
            "# My Document\n\nContent.\n\n---\n\n## Comments\n\n### Old Author (2025-01-01)\nOld comment.\n",
            encoding="utf-8",
        )
        docx_path = _make_docx_with_comments(tmp_path)

        append_comments_to_markdown(md_path, docx_path)

        content = md_path.read_text(encoding="utf-8")
        assert "Old Author" not in content
        assert "Old comment" not in content
        assert "Jane Smith" in content
        assert content.count("## Comments") == 1

    def test_no_comments_returns_false(self, tmp_path):
        """Returns False when .docx has no comments."""
        md_path = tmp_path / "doc.md"
        md_path.write_text("# Doc\nContent.\n", encoding="utf-8")
        docx_path = _make_simple_docx(tmp_path)

        result = append_comments_to_markdown(md_path, docx_path)
        assert result is False

        content = md_path.read_text(encoding="utf-8")
        assert "## Comments" not in content

    def test_preserves_body_content(self, tmp_path):
        """Body content is fully preserved when appending comments."""
        body = "# Title\n\n## Section 1\n\nParagraph one.\n\n## Section 2\n\nParagraph two.\n"
        md_path = tmp_path / "doc.md"
        md_path.write_text(body, encoding="utf-8")
        docx_path = _make_docx_with_comments(tmp_path)

        append_comments_to_markdown(md_path, docx_path)

        content = md_path.read_text(encoding="utf-8")
        assert "# Title" in content
        assert "## Section 1" in content
        assert "Paragraph one." in content
        assert "## Section 2" in content
        assert "Paragraph two." in content


# ---------------------------------------------------------------------------
# Tests: KNOWLEDGE_EXTENSIONS whitelist
# ---------------------------------------------------------------------------

class TestKnowledgeExtensions:
    def test_includes_text_formats(self):
        from brain_sync.fileops import TEXT_EXTENSIONS
        assert ".md" in TEXT_EXTENSIONS
        assert ".txt" in TEXT_EXTENSIONS
        assert ".csv" in TEXT_EXTENSIONS
        assert ".json" in TEXT_EXTENSIONS

    def test_includes_image_formats(self):
        from brain_sync.fileops import IMAGE_EXTENSIONS
        assert ".png" in IMAGE_EXTENSIONS
        assert ".jpg" in IMAGE_EXTENSIONS
        assert ".jpeg" in IMAGE_EXTENSIONS

    def test_excludes_binary_formats(self):
        from brain_sync.fileops import KNOWLEDGE_EXTENSIONS
        assert ".pdf" not in KNOWLEDGE_EXTENSIONS
        assert ".docx" not in KNOWLEDGE_EXTENSIONS
        assert ".zip" not in KNOWLEDGE_EXTENSIONS
        assert ".exe" not in KNOWLEDGE_EXTENSIONS

    def test_union(self):
        from brain_sync.fileops import IMAGE_EXTENSIONS, KNOWLEDGE_EXTENSIONS, TEXT_EXTENSIONS
        assert KNOWLEDGE_EXTENSIONS == TEXT_EXTENSIONS | IMAGE_EXTENSIONS
